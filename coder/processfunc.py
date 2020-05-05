import pyqrcode
import re
import os
from math import ceil
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

def getFileName(pathToFile):
	name = re.search(r'\_(.*?)\_\.(docx|doc)', pathToFile)
	return name.group(0)
	
def downloadFileOnServer(file):
	pathToUploadedFile = os.path.join(BASE_DIR,'coder/temp/' + file.name)
	with open(pathToUploadedFile, 'wb+') as destination:
		for chunk in file.chunks():
			destination.write(chunk)
	return pathToUploadedFile  

def processDonloadedFile(filename):
    document = Document(filename)
    text = re.sub(r'\s|\t|\n', '', readTextFrowWordFile(document))
    fileToSave = getQrCodeFromText(text, filename)
    return addQrCodeToWordFile(document, filename, fileToSave)
    
def readTextFrowWordFile(document):
    return ''.join([paragraph.text for paragraph in document.paragraphs])

def getQrCodeFromText(text, filePath):
    lengthText = len(text)
    skip = calcSkipCount(lengthText)
    fileToSave = filePath + '.png'
    big_code = pyqrcode.create(text[::skip], error='L', version=40, encoding='utf-8')
    big_code.png(fileToSave, scale=6)
    return fileToSave

def calcSkipCount(lettersCount, maxLength=1630):
    return ceil(lettersCount / maxLength)

def addQrCodeToWordFile(document, wordFilePath, qrCodePng):
    paragraphPic = document.add_paragraph()
    paragraphPic.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraphPic.add_run()
    run.add_picture(qrCodePng, width=Inches(3))
    wordFilePathNew = wordFilePath.replace('.','_.')
    document.save(wordFilePathNew)
    removeTempFiles([qrCodePng, wordFilePath])
    return wordFilePathNew
    
def removeTempFiles(files):
    [os.remove(file) for file in files]

def isMsWordFile(filePath):
    return filePath.endswith('docx') or filePath.endswith('doc')
